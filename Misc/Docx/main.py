# writing to docx file
def writing():
    from docx import Document
    document = Document()

    document.add_heading("Hello World", 0)
    # 0 - title
    # 1 - heading 1
    # 2 - heading 2

    p = document.add_paragraph("This is a sample text.")
    document.add_picture("image_1.png")
    p.add_run("\nThis text is bold").bold = True # .italic

    document.add_paragraph("This is item One", style="List Bullet")
    document.add_paragraph("This is item Two", style="List Bullet")
    document.add_paragraph("This is item Three", style="List Bullet")
    document.add_paragraph("This is item Four", style="List Bullet")

    table_header = ["Name", "Age", "Job"]

    some_data = [
        ["John", 46, "programmer"],
        ["Mary", 55, "Chef"],
        ["Bob", 23, "Accountant"],
    ]

    table = document.add_table(rows=1, cols=3)
    for i in range(len(some_data)):
        table.rows[0].cells[i].text = table_header[i]

    for name, age, job in some_data:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = str(age)
        cells[2].text = job


    document.save("test.docx")

# reading a docx file
def reading():
    from docx.api import Document

    doc = Document("Operational Guide_Railway Mail Service _Version 3.0 Dated 08.03.2025.docx")

    # for p in doc.paragraphs:
    #     if p.style.name.startswith("Heading"):
    #         print(p.text)

    # for table in doc.tables:
    #     print("+"*80)
    #     for row in table.rows:
    #         print("|".join([cell.text for cell in row.cells]))

    sections = doc.sections
    print(f"Total sections found: {len(sections)}")

    

# writing()
reading()


